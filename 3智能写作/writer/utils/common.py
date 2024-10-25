from pydantic import BaseModel
from typing import List, Union, Any, Callable, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import re
import os
import time
from datetime import datetime
from docx import Document
 
def colored_decorator(color):   
    def print_colored(func):  
        async def wrapper(self, *args, **kwargs):
                # ANSI escape codes for colored terminal output (optional, can be removed if not needed)
                print(color)
                contexts = await func(self, *args, **kwargs)  
                print("\033[0m")
                return contexts             
        return wrapper
    return print_colored


def print_time(func):  
    async def wrapper(self, *args, **kwargs):
            # ANSI escape codes for colored terminal output (optional, can be removed if not needed)
            start_time = datetime.now()
            objs = await func(self, *args, **kwargs)  
            end_time = datetime.now()
            print(f"执行时间：{(end_time - start_time).total_seconds()}秒")
            return objs             
    return wrapper


class WriteOutFile:
    
    @staticmethod
    def post_processes(title: str, context: str) -> str:
            """Post-processes the context by removing the first line that match the title."""
            normalize_text = lambda x:  re.sub(r'[^\u4e00-\u9fa5]+', '', x)
            split_context = context.split('\n') if context else []
            prefix = split_context[0] if split_context else ''
            if  normalize_text(title) == normalize_text(prefix): 
                context = '\n'.join(split_context[1:])
            return context  
    

    def write_markdown_file(self, topic: str, tasks: Any, output_path: Union[str, Path]):
        document = ''
        for title, content, level  in tasks:
            document += f"{'#' * level} {title}\n\n"
            # Post-process and add it to the document
            content = self.post_processes(title, content)
            document +=  f"{content}\n"
        # Save the document
        output_path = Path(output_path)
        if  not output_path.exists():
                os.makedirs(output_path)
        try:
            with open(output_path / f'{topic}.md',  'w', encoding='utf-8') as f:
                f.write(document)
        except Exception as e:
            print(f"An error occurred while saving the document: {e}")  
    
    
    def write_word_file(self, topic: str, tasks: Any, output_path: Union[str, Path]):
        """
        Writes tasks to a Word document.

        topic (str): The main topic of the document.
        tasks (List[Tuple[str, str, int]]): A list of tuples containing the title, content, and heading level of each task.
        output_path (Union[str, Path]): The file path where the document will be saved.
        """
        # Ensure the write_out_file is a Path object
        document = Document()
        document.add_heading(topic, level=0)
        # Process each task
        for title, content, level  in tasks:
            # Add a heading for the task
            document.add_heading(title, level=level)
            # Post-process and add it to the document
            content = self.post_processes(title, content)
            document.add_paragraph(content)
        document.add_page_break()
        # Save the document
        output_path = Path(output_path)
        if  not output_path.exists():
                os.makedirs(output_path)
        try:
            document.save(output_path / f'{topic}.docx')
        except Exception as e:
            print(f"An error occurred while saving the document: {e}")   
    
  
       
         